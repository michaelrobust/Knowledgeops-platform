"""
Document Parser Service
Handles PDF, DOCX, and TXT file parsing and text extraction
"""

import os
import logging
from typing import List, Dict, Any
from pathlib import Path

# PDF processing
try:
    import pdfplumber
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️ PDF processing libraries not available")

# DOCX processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ DOCX processing library not available")

logger = logging.getLogger(__name__)

class DocumentParser:
    """Document parser for various file formats"""
    
    def __init__(self):
        self.supported_types = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.txt': self._parse_txt
        }
    
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """
        Parse document and extract text content
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Get file extension
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext not in self.supported_types:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Parse the document
            parser_func = self.supported_types[file_ext]
            result = parser_func(file_path)
            
            # Add common metadata
            file_stat = os.stat(file_path)
            result.update({
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'file_size': file_stat.st_size,
                'file_type': file_ext,
                'word_count': len(result.get('content', '').split()),
                'char_count': len(result.get('content', ''))
            })
            
            logger.info(f"Successfully parsed {file_path}: {result['word_count']} words")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {str(e)}")
            raise
    
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Parse PDF file using pdfplumber and PyMuPDF"""
        if not PDF_AVAILABLE:
            raise ImportError("PDF processing libraries not installed")
        
        content = ""
        pages_info = []
        
        try:
            # Try pdfplumber first (better for text extraction)
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text() or ""
                    content += page_text + "\n\n"
                    
                    pages_info.append({
                        'page_number': page_num,
                        'text_length': len(page_text),
                        'has_text': bool(page_text.strip())
                    })
        
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyMuPDF: {str(e)}")
            
            # Fallback to PyMuPDF
            try:
                doc = fitz.open(file_path)
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    page_text = page.get_text()
                    content += page_text + "\n\n"
                    
                    pages_info.append({
                        'page_number': page_num + 1,
                        'text_length': len(page_text),
                        'has_text': bool(page_text.strip())
                    })
                doc.close()
            
            except Exception as fallback_error:
                raise Exception(f"Both PDF parsers failed: {str(e)}, {str(fallback_error)}")
        
        return {
            'content': content.strip(),
            'pages': pages_info,
            'total_pages': len(pages_info),
            'extraction_method': 'pdfplumber/pymupdf'
        }
    
    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Parse DOCX file using python-docx"""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX processing library not installed")
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            content = ""
            
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                if para_text:  # Skip empty paragraphs
                    paragraphs.append({
                        'paragraph_number': i + 1,
                        'text': para_text,
                        'length': len(para_text)
                    })
                    content += para_text + "\n\n"
            
            # Extract text from tables (if any)
            tables_content = []
            for table_num, table in enumerate(doc.tables):
                table_text = ""
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    table_text += row_text + "\n"
                
                if table_text.strip():
                    tables_content.append({
                        'table_number': table_num + 1,
                        'content': table_text.strip()
                    })
                    content += f"\n[Table {table_num + 1}]\n{table_text}\n"
            
            return {
                'content': content.strip(),
                'paragraphs': paragraphs,
                'tables': tables_content,
                'total_paragraphs': len(paragraphs),
                'total_tables': len(tables_content),
                'extraction_method': 'python-docx'
            }
            
        except Exception as e:
            raise Exception(f"Failed to parse DOCX file: {str(e)}")
    
    def _parse_txt(self, file_path: str) -> Dict[str, Any]:
        """Parse TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin1', 'cp1252']
            content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        used_encoding = encoding
                        break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise Exception("Could not decode file with any supported encoding")
            
            # Split into lines for analysis
            lines = content.split('\n')
            non_empty_lines = [line.strip() for line in lines if line.strip()]
            
            return {
                'content': content.strip(),
                'total_lines': len(lines),
                'non_empty_lines': len(non_empty_lines),
                'encoding': used_encoding,
                'extraction_method': 'text_file'
            }
            
        except Exception as e:
            raise Exception(f"Failed to parse TXT file: {str(e)}")
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
        """
        Split text into smaller chunks for processing
        
        Args:
            text: Input text to chunk
            chunk_size: Maximum size of each chunk
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks with metadata
        """
        if not text or not text.strip():
            return []
        
        chunks = []
        start = 0
        chunk_id = 1
        
        while start < len(text):
            # Calculate end position
            end = start + chunk_size
            
            # If we're not at the end, try to break at word boundary
            if end < len(text):
                # Look backwards for the last space within reasonable distance
                word_boundary = text.rfind(' ', start, end)
                if word_boundary > start + chunk_size // 2:  # Only if reasonable
                    end = word_boundary
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunks.append({
                    'chunk_id': chunk_id,
                    'text': chunk_text,
                    'start_pos': start,
                    'end_pos': end,
                    'length': len(chunk_text),
                    'word_count': len(chunk_text.split())
                })
                chunk_id += 1
            
            # Move start position (with overlap)
            start = end - overlap if end < len(text) else len(text)
        
        logger.info(f"Created {len(chunks)} chunks from text of {len(text)} characters")
        return chunks

# Convenience function for direct use
def parse_document(file_path: str) -> Dict[str, Any]:
    """Parse a document file and return extracted content"""
    parser = DocumentParser()
    return parser.parse_document(file_path)

def chunk_document(file_path: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict[str, Any]]:
    """Parse a document and return it as chunks"""
    parser = DocumentParser()
    result = parser.parse_document(file_path)
    content = result.get('content', '')
    chunks = parser.chunk_text(content, chunk_size, overlap)
    
    # Add document metadata to each chunk
    for chunk in chunks:
        chunk.update({
            'source_file': result.get('file_name'),
            'source_path': result.get('file_path'),
            'file_type': result.get('file_type'),
            'extraction_method': result.get('extraction_method')
        })
    
    return chunks
